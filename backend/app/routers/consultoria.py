"""
Consultoría (honorarios / fees) — módulo sencillo.

Un **contrato** con pocos datos (cliente=Productor, fecha inicio, duración en meses [o indefinido],
frecuencia de cobro, importe por cobro, sujeto a impuestos + % IVA). De cada cobro se **genera un
recibo tipo 'Consultoría'** cuando toca (no se crean los futuros). El recibo reutiliza el modelo
`recibos`: Base Imponible = `comision_retenida` = importe; IVA = `impuestos_recibo` (lo que ya usa
el cierre contable). El cobro/edición se gestiona luego en la pantalla de Recibos como cualquier otro.
"""
from __future__ import annotations

import calendar
import datetime as dt
import os
import tempfile
import zipfile
from decimal import Decimal

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, ConfigDict
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from ..config import settings
from ..db import get_db
from ..models.maestras import ConsultoriaContrato, CuentaBancaria, Productor, Recibo
from .recibos import _exigir_mes_abierto, _recompute, _siguiente_numero

router = APIRouter(tags=["Consultoría"])

PASO_MESES = {"Mensual": 1, "Trimestral": 3, "Semestral": 6, "Anual": 12}
MESES_ES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
            "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]


def _add_months(d: dt.date, n: int) -> dt.date:
    m = d.month - 1 + n
    y, mo = d.year + m // 12, m % 12 + 1
    return dt.date(y, mo, min(d.day, calendar.monthrange(y, mo)[1]))


def _fechas_cobro(c: ConsultoriaContrato) -> list[dt.date]:
    """Fechas de los cobros del contrato según su frecuencia y duración. Indefinido (sin duración):
    hasta hoy + un periodo por delante (siempre hay un cobro 'actual/próximo' que emitir)."""
    if c.frecuencia == "Único":
        return [c.fecha_inicio]
    paso = PASO_MESES.get(c.frecuencia)
    if not paso:
        return [c.fecha_inicio]
    fin = _add_months(c.fecha_inicio, c.duracion_meses) if c.duracion_meses else None
    tope = fin if fin is not None else _add_months(dt.date.today(), paso)
    fechas, k = [], 0
    while k < 1200:
        f = _add_months(c.fecha_inicio, k * paso)
        if (fin is not None and f >= fin) or (fin is None and f > tope):
            break
        fechas.append(f)
        k += 1
    return fechas


def _iva(c: ConsultoriaContrato, base: Decimal) -> Decimal:
    if not c.sujeto_impuestos:
        return Decimal("0.00")
    return (base * (c.impuestos_porc or Decimal(0)) / Decimal(100)).quantize(Decimal("0.01"))


def _cobro_debido(c: ConsultoriaContrato, f: dt.date, hoy: dt.date, generado: bool) -> bool:
    """Un cobro 'cuenta' y se muestra si ya tiene recibo o su AVISO ya ha saltado: `aviso_dias_antes`
    antes de la fecha de facturación (según el contrato). Los cobros futuros aún sin aviso no salen."""
    if generado:
        return True
    return _fecha_facturacion(c, f) - dt.timedelta(days=int(c.aviso_dias_antes or 0)) <= hoy


# ── Schemas ──
class ContratoIn(BaseModel):
    productor_id: int
    concepto: str | None = None
    fecha_inicio: dt.date
    duracion_meses: int | None = None      # None = indefinido
    frecuencia: str                        # Mensual/Trimestral/Semestral/Anual/Único
    importe: Decimal
    sujeto_impuestos: bool = True
    impuestos_porc: Decimal = Decimal("21")
    moneda: str = "EUR"
    cuenta_bancaria_id: int | None = None
    dia_facturacion: int | None = None     # día del mes en que se factura (None = día de fecha_inicio)
    aviso_dias_antes: int = 5
    estado: str = "Activo"
    notas: str | None = None


class ContratoUpdate(BaseModel):
    productor_id: int | None = None
    concepto: str | None = None
    fecha_inicio: dt.date | None = None
    duracion_meses: int | None = None
    frecuencia: str | None = None
    importe: Decimal | None = None
    sujeto_impuestos: bool | None = None
    impuestos_porc: Decimal | None = None
    moneda: str | None = None
    cuenta_bancaria_id: int | None = None
    dia_facturacion: int | None = None
    aviso_dias_antes: int | None = None
    estado: str | None = None
    notas: str | None = None


class ContratoRead(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    productor_id: int
    productor_nombre: str | None = None
    concepto: str | None = None
    fecha_inicio: dt.date
    duracion_meses: int | None = None
    frecuencia: str
    importe: Decimal
    sujeto_impuestos: bool
    impuestos_porc: Decimal
    moneda: str
    cuenta_bancaria_id: int | None = None
    cuenta_bancaria_nombre: str | None = None
    dia_facturacion: int | None = None
    aviso_dias_antes: int = 5
    estado: str
    notas: str | None = None
    n_cobros: int = 0          # cobros previstos
    n_generados: int = 0       # recibos ya generados
    proximo_cobro: dt.date | None = None  # primer cobro sin recibo


def _serializar(db: Session, c: ConsultoriaContrato) -> ContratoRead:
    d = ContratoRead.model_validate(c)
    d.productor_nombre = c.productor.nombre if c.productor else None
    d.cuenta_bancaria_nombre = c.cuenta_bancaria.nombre if c.cuenta_bancaria else None
    fechas = _fechas_cobro(c)
    generados = {r.periodo for r in db.scalars(
        select(Recibo).where(Recibo.consultoria_id == c.id)
    ).all()}
    hoy = dt.date.today()
    # Solo cuentan los cobros DEBIDOS: ya generados o con su aviso ya saltado (no los futuros).
    debidos = [f for f in fechas if _cobro_debido(c, f, hoy, f.strftime("%Y-%m") in generados)]
    d.n_cobros = len(debidos)
    d.n_generados = len(generados)
    # Próximo cobro PENDIENTE de hacer: el primer cobro debido sin recibo (None si nada toca aún).
    d.proximo_cobro = next(
        (f for f in fechas if f.strftime("%Y-%m") not in generados and _cobro_debido(c, f, hoy, False)),
        None,
    )
    return d


@router.get("/consultoria")
def listar(db: Session = Depends(get_db)):
    cs = db.scalars(select(ConsultoriaContrato).order_by(ConsultoriaContrato.id.desc())).all()
    return [_serializar(db, c) for c in cs]


@router.get("/consultoria/{contrato_id}")
def obtener(contrato_id: int, db: Session = Depends(get_db)):
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    return _serializar(db, c)


@router.post("/consultoria", status_code=201)
def crear(payload: ContratoIn, db: Session = Depends(get_db)):
    c = ConsultoriaContrato(**payload.model_dump())
    db.add(c)
    db.commit()
    db.refresh(c)
    return _serializar(db, c)


@router.put("/consultoria/{contrato_id}")
def editar(contrato_id: int, payload: ContratoUpdate, db: Session = Depends(get_db)):
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    for k, v in payload.model_dump(exclude_unset=True).items():
        setattr(c, k, v)
    db.commit()
    db.refresh(c)
    return _serializar(db, c)


@router.delete("/consultoria/{contrato_id}", status_code=204)
def borrar(contrato_id: int, db: Session = Depends(get_db)):
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    db.delete(c)  # los recibos enlazados quedan con consultoria_id NULL (SET NULL)
    db.commit()


@router.get("/consultoria/{contrato_id}/cobros")
def cobros(contrato_id: int, db: Session = Depends(get_db)):
    """Calendario de cobros del contrato, con su recibo si ya se generó."""
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    recibos = {r.periodo: r for r in db.scalars(select(Recibo).where(Recibo.consultoria_id == c.id)).all()}
    hoy = dt.date.today()
    out = []
    for f in _fechas_cobro(c):
        per = f.strftime("%Y-%m")
        r = recibos.get(per)
        # Solo se muestra un cobro cuando toca: ya generado o con su aviso ya saltado (no los futuros).
        if not _cobro_debido(c, f, hoy, r is not None):
            continue
        base = Decimal(c.importe or 0)
        iva = _iva(c, base)
        adeu = r.prima_adeudada if r else None
        cobrado = bool(r and adeu and (r.prima_cobrada or 0) >= adeu)
        out.append({
            "periodo": per, "fecha": f.isoformat(),
            "base": float(base), "iva": float(iva), "total": float(base + iva),
            "recibo_id": r.id if r else None, "recibo_numero": r.numero if r else None,
            "recibo_cobrado": cobrado,
        })
    return {"contrato_id": c.id, "moneda": c.moneda, "cobros": out}


class GenerarCobro(BaseModel):
    periodo: str   # 'YYYY-MM' del cobro a generar


def _crear_recibo_cobro(db: Session, c: ConsultoriaContrato, fecha: dt.date, periodo: str) -> Recibo:
    """Crea el Recibo tipo 'Consultoría' del cobro (no comitea). Falla si ya existe o el mes está
    cerrado."""
    ya = db.scalar(select(Recibo).where(Recibo.consultoria_id == c.id, Recibo.periodo == periodo))
    if ya is not None:
        raise HTTPException(status_code=409, detail=f"El cobro {periodo} ya tiene recibo ({ya.numero}).")
    _exigir_mes_abierto(db, fecha)
    base = Decimal(c.importe or 0)
    iva = _iva(c, base)
    cuenta = c.cuenta_bancaria.nombre if c.cuenta_bancaria else None
    r = Recibo(
        consultoria_id=c.id, periodo=periodo, anio=fecha.year, estado="Emitido",
        numero=_siguiente_numero(db, fecha.year),
        tipo_poliza="Consultoría", asegurado=(c.productor.nombre if c.productor else None),
        corredor=((c.productor.alias or c.productor.nombre) if c.productor else None),
        pagador=(c.productor.nombre if c.productor else None),
        ramo="Consultoría", moneda=c.moneda, cuenta=cuenta,
        fecha_efecto=fecha, fecha_vencimiento=fecha, fecha_contable=fecha,
        # El listado de recibos usa fecha_efecto_recibo / fecha_vcto_recibo (no fecha_efecto).
        fecha_efecto_recibo=fecha, fecha_vcto_recibo=fecha,
        honorarios=base, comision_retenida=base, impuestos_porc=c.impuestos_porc,
        impuestos_recibo=iva, prima_bruta_recibo=base + iva, prima_adeudada=base + iva,
    )
    _recompute(r)
    db.add(r)
    return r


@router.post("/consultoria/{contrato_id}/cobros/generar", status_code=201)
def generar_cobro(contrato_id: int, payload: GenerarCobro, db: Session = Depends(get_db)):
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    fecha = next((f for f in _fechas_cobro(c) if f.strftime("%Y-%m") == payload.periodo), None)
    if fecha is None:
        raise HTTPException(status_code=422, detail=f"El periodo {payload.periodo} no es un cobro de este contrato.")
    r = _crear_recibo_cobro(db, c, fecha, payload.periodo)
    db.commit()
    db.refresh(r)
    base = Decimal(c.importe or 0)
    return {"recibo_id": r.id, "numero": r.numero, "periodo": r.periodo, "total": float(base + _iva(c, base))}


# ──────────────────────────── Factura Word ────────────────────────────
def _num_es(x: Decimal) -> str:
    """123456.7 -> '123.456,70' (formato es-ES)."""
    s = f"{Decimal(x):,.2f}"                    # '123,456.70'
    return s.replace(",", "·").replace(".", ",").replace("·", ".")


def _fecha_facturacion(c: ConsultoriaContrato, f: dt.date) -> dt.date:
    dia = c.dia_facturacion or c.fecha_inicio.day
    return dt.date(f.year, f.month, min(dia, calendar.monthrange(f.year, f.month)[1]))


def _nombre_corto(nombre: str | None) -> str:
    """'Insurart, S.L.' -> 'Insurart' (para nombre de archivo/carpeta)."""
    n = (nombre or "Cliente").split(",")[0].strip()
    return "".join(ch for ch in n if ch not in '\\/:*?"<>|').strip() or "Cliente"


def _set_token(p, token: str, valor: str, ultimo: bool = False) -> bool:
    """Sustituye en el párrafo el run cuyo texto (sin espacios) sea exactamente `token` por `valor`,
    preservando el formato. `ultimo`=True usa la última coincidencia (cuando etiqueta y valor
    comparten texto, p. ej. 'Cliente')."""
    idxs = [i for i, r in enumerate(p.runs) if r.text.strip() == token]
    if not idxs:
        return False
    p.runs[idxs[-1] if ultimo else idxs[0]].text = valor
    return True


def _generar_factura_docx(c: ConsultoriaContrato, r: Recibo, fecha_fact: dt.date,
                          pago_n: int, pago_t: int, cta: CuentaBancaria | None) -> str:
    """Rellena la plantilla de factura para el recibo `r` y la guarda en
    <facturas_dir>\\<año>\\Facturas Emitidas\\<Cliente>\\<numero> <Cliente> <Mes>.docx."""
    import docx  # carga perezosa

    plantilla = settings.factura_plantilla
    if not os.path.isfile(plantilla):
        raise HTTPException(status_code=502, detail=f"No se encuentra la plantilla de factura: {plantilla}")

    # La plantilla es .dotx → se convierte a .docx (cambiando el content-type) en un temporal.
    tmp = os.path.join(tempfile.gettempdir(), "_factura_plantilla.docx")
    with zipfile.ZipFile(plantilla) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zo:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == "[Content_Types].xml":
                data = data.replace(b"wordprocessingml.template.main+xml", b"wordprocessingml.document.main+xml")
            zo.writestr(it, data)

    prod = c.productor
    base = Decimal(c.importe or 0)
    iva = _iva(c, base)
    moneda = c.moneda or "EUR"
    concepto = f"Servicios Profesionales de seguros {MESES_ES[fecha_fact.month].lower()} {fecha_fact.year}"

    d = docx.Document(tmp)
    for p in d.paragraphs:
        t = p.text.strip()
        if "NumeroRecibo" in t:
            _set_token(p, "NumeroRecibo", r.numero or "")
        elif t.startswith("Cliente"):
            _set_token(p, "Cliente", prod.nombre if prod else "", ultimo=True)
        elif t.startswith("CIFCliente"):
            _set_token(p, "CIFCliente", (prod.cif if prod else "") or "")
        elif t.startswith("Fecha"):
            _set_token(p, "FechaFactura", fecha_fact.strftime("%d.%m.%Y"))
        elif "Pagos" in t:                              # 'Número de Pagos: <frecuencia>'
            _set_token(p, "Pago", c.frecuencia or "")
        elif t.startswith("Pago"):                      # 'Pago: <n> de <total>'
            _set_token(p, "Recibo", str(pago_n))
            _set_token(p, "RecibosTotales", str(pago_t))
        elif t.startswith("Moneda"):
            _set_token(p, "Moneda", moneda)
        elif t.startswith("Concepto"):
            _set_token(p, "Concepto", concepto)

    # Tabla de importes.
    ti = d.tables[0]
    _set_token(ti.rows[0].cells[2].paragraphs[0], "Importe", _num_es(base))
    _set_token(ti.rows[1].cells[1].paragraphs[0], "ImpuestosPorc", f"{_num_es(c.impuestos_porc or 0)}%")
    _set_token(ti.rows[1].cells[2].paragraphs[0], "ImpuestosRecibo", _num_es(iva))
    _set_token(ti.rows[2].cells[2].paragraphs[0], "ImporteTotal", _num_es(base + iva))

    # Tabla de banco (si hay cuenta). IBAN agrupado en bloques de 4; oficina = posiciones 9-12.
    if cta is not None and len(d.tables) > 1:
        iban = (cta.iban or "").replace(" ", "")
        iban_fmt = " ".join(iban[i:i + 4] for i in range(0, len(iban), 4))
        oficina = iban[8:12] if len(iban) >= 12 else ""
        for row in d.tables[1].rows:
            cell = row.cells[0]
            for p in cell.paragraphs:
                _set_token(p, "Banco", cta.banco or "")
                _set_token(p, "Oficina", oficina, ultimo=True)
                _set_token(p, "BIC", (cta.swift_bic or ""))
                _set_token(p, "Cuenta", iban_fmt)

    corto = _nombre_corto(prod.nombre if prod else None)
    carpeta = os.path.join(settings.facturas_dir, str(fecha_fact.year), "Facturas Emitidas", corto)
    os.makedirs(carpeta, exist_ok=True)
    destino = os.path.join(carpeta, f"{r.numero} {corto} {MESES_ES[fecha_fact.month]}.docx")
    d.save(destino)
    return destino


@router.post("/consultoria/{contrato_id}/cobros/generar-factura", status_code=201)
def generar_factura(contrato_id: int, payload: GenerarCobro, db: Session = Depends(get_db)):
    """Genera (si falta) el recibo del cobro y produce el Word de la factura, listo para enviar.
    Devuelve la ruta del documento."""
    c = db.get(ConsultoriaContrato, contrato_id)
    if c is None:
        raise HTTPException(status_code=404, detail=f"Contrato {contrato_id} no encontrado")
    fechas = _fechas_cobro(c)
    fecha = next((f for f in fechas if f.strftime("%Y-%m") == payload.periodo), None)
    if fecha is None:
        raise HTTPException(status_code=422, detail=f"El periodo {payload.periodo} no es un cobro de este contrato.")

    r = db.scalar(select(Recibo).where(Recibo.consultoria_id == c.id, Recibo.periodo == payload.periodo))
    if r is None:
        r = _crear_recibo_cobro(db, c, fecha, payload.periodo)
        db.commit()
        db.refresh(r)

    # Cuenta para el bloque bancario: la del contrato; si no, una de Gastos (honorarios) activa.
    cta = c.cuenta_bancaria
    if cta is None:
        cta = db.scalar(
            select(CuentaBancaria).where(CuentaBancaria.activa.is_(True), CuentaBancaria.categoria == "Gastos")
            .order_by(CuentaBancaria.id)
        )

    fecha_fact = _fecha_facturacion(c, fecha)
    pago_n = next((i + 1 for i, f in enumerate(fechas) if f.strftime("%Y-%m") == payload.periodo), 1)
    pago_t = len(fechas)
    ruta = _generar_factura_docx(c, r, fecha_fact, pago_n, pago_t, cta)
    return {"recibo_id": r.id, "numero": r.numero, "periodo": r.periodo, "archivo": ruta}
