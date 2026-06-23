"""
LPAN (London Premium Advice Note) y FDO por binder.

Flujo Lloyd's/Xchanging:
  - Por cada (binder, risk code) se genera un FDO (Declaración) que se envía a Xchanging; éste
    devuelve un `signing_number`.
  - A partir de ahí, cada periodo de Premium BDX genera un LPAN (nota de pago) que agrupa las
    líneas de ese risk code en el periodo y cuelga del signing del FDO.

Importes del LPAN (sobre las líneas del Premium, incluidas en premium):
  - gross_premium = Σ total_gwp_our_line            (campo 18, GWP our line)
  - brokerage     = Σ (commission_coverholder + brokerage_amount)   (campo 19)
  - tax           = Σ total_taxes_levies            (campo 17, IPT)
  - net_premium   = Σ final_net_premium_uw          (campo 25, Bureau NA Premium)

Solo se genera un LPAN si todas las líneas del grupo están cobradas.
"""
from __future__ import annotations

import datetime as dt
import os
import tempfile
import zipfile
from decimal import Decimal

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.orm import Session, load_only, selectinload

from ..config import settings
from ..db import get_db
from ..models.maestras import Bdx, BdxLinea, Binder, BinderSeccion, Fdo, Lpan, Poliza, SeccionRiskCode

router = APIRouter(tags=["LPAN"])

D0 = Decimal("0")
MESES_ES = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
            "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]


def _d(v) -> Decimal:
    return v if isinstance(v, Decimal) else Decimal(str(v or 0))


def _umr_part(agreement: str | None) -> str:
    """Parte del UMR para el Broker Reference: el agreement sin el código de 3 letras del
    coverholder (p.ej. 'CY0926ALE' -> 'CY0926', 'PI2926CRO' -> 'PI2926')."""
    a = (agreement or "").strip()
    return a[:-3] if len(a) > 3 and a[-3:].isalpha() else a


def _broker_ref(agreement: str | None, section: int, risk_code: str) -> str:
    """Nombre del FDO: '<parte del UMR> FDO-S<sección>-<risk code>' (p.ej. 'CY0926 FDO-S1-PC')."""
    return f"{_umr_part(agreement)} FDO-S{section}-{risk_code}"


def _set_celda(cell, valor: str) -> None:
    """Pone `valor` en una celda conservando el formato del primer run."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = valor
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(valor)


def _generar_fdo_docx(carpeta: str, broker_ref: str, ref1: str, umr: str | None, signing: str | None) -> str:
    """Copia la plantilla LPAN (formulario de tokens) y la rellena para un FDO, guardándola como
    '<broker_ref>.docx' en `carpeta`. `ref1` = Broker Reference 1 (parte del UMR, campo 10);
    `broker_ref` va en Broker Reference 2 (campo 11). Devuelve la ruta del documento."""
    import docx  # carga perezosa: solo al generar

    plantilla = settings.lpan_plantilla
    if not os.path.isfile(plantilla):
        raise HTTPException(status_code=502, detail=f"No se encuentra la plantilla LPAN: {plantilla}")
    if not carpeta or not os.path.isdir(carpeta):
        raise HTTPException(status_code=400, detail=f"La carpeta indicada no existe: {carpeta!r}")

    # La plantilla es .dotx; se convierte a .docx (cambiando el content-type) en un temporal.
    tmp = os.path.join(tempfile.gettempdir(), "_fdo_plantilla.docx")
    with zipfile.ZipFile(plantilla) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zo:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == "[Content_Types].xml":
                data = data.replace(b"wordprocessingml.template.main+xml", b"wordprocessingml.document.main+xml")
            zo.writestr(it, data)

    d = docx.Document(tmp)
    # Tokens que se rellenan SIEMPRE (mismo valor en todas sus apariciones).
    todos = {"Premium": "FDO", "Yes/No": "No"}
    # Tokens de "línea" (la plantilla tiene 3 filas): solo la 1ª lleva valor, las demás se vacían.
    una_vez = {
        "Bureau": signing or "", "BrokerRef1": ref1, "BrokerRef2": broker_ref,
        "Line": "", "Taxes": "", "GrossPremium": "", "Brokerage": "",
        "OCurrency": "FDO", "SCurrency": "FDO", "BureauPremium": "", "UMR": umr or "",
    }
    vistos_tok: set[str] = set()
    vistos_tc: set = set()
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell._tc in vistos_tc:
                    continue
                vistos_tc.add(cell._tc)
                k = cell.text.strip()
                if k in todos:
                    _set_celda(cell, todos[k])
                elif k in una_vez:
                    _set_celda(cell, una_vez[k] if k not in vistos_tok else "")
                    vistos_tok.add(k)

    destino = os.path.join(carpeta, f"{broker_ref}.docx")
    d.save(destino)
    return destino


def _nombre_lpan(agreement: str | None, periodo: str, section: int, risk_code: str) -> str:
    """Nombre/Broker Ref 2 del LPAN (patrón histórico): '<UMR> <MM> BDX-S<sec>-<rc>-<MMAA>'
    (p.ej. periodo 2019-07, PI0219CRO, S1, E7 -> 'PI0219 07 BDX-S1-E7-0719')."""
    yyyy, mm = (periodo.split("-") + ["", ""])[:2]
    return f"{_umr_part(agreement)} {mm} BDX-S{section}-{risk_code}-{mm}{yyyy[2:]}"


def _num_lpan(v) -> str:
    """Importe para el documento (formato inglés del LPAN: 1.234,56 -> '1,234.56')."""
    return f"{_d(v):,.2f}"


def _pct_lpan(v) -> str:
    """Porcentaje para el documento, sin ceros sobrantes: 5 -> '5%', 32.5 -> '32.5%'."""
    s = f"{_d(v):.2f}".rstrip("0").rstrip(".")
    return f"{s}%"


def _generar_lpan_docx(carpeta: str, nombre: str, signing: str | None, broker_ref1: str,
                       umr: str | None, gross, brokerage, tax, net, moneda: str) -> str:
    """Rellena la plantilla LPAN con las cifras reales del bloque y la guarda como '<nombre>.docx'
    en `carpeta`. Devuelve la ruta del documento. `broker_ref1` (casilla 10) = código completo del
    agreement, con las 3 letras del coverholder al final."""
    import docx  # carga perezosa

    plantilla = settings.lpan_plantilla
    if not os.path.isfile(plantilla):
        raise HTTPException(status_code=502, detail=f"No se encuentra la plantilla LPAN: {plantilla}")
    if not carpeta or not os.path.isdir(carpeta):
        raise HTTPException(status_code=400, detail=f"La carpeta indicada no existe: {carpeta!r}")

    # La plantilla es .dotx; se convierte a .docx (cambiando el content-type) en un temporal.
    tmp = os.path.join(tempfile.gettempdir(), "_lpan_plantilla.docx")
    with zipfile.ZipFile(plantilla) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zo:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == "[Content_Types].xml":
                data = data.replace(b"wordprocessingml.template.main+xml", b"wordprocessingml.document.main+xml")
            zo.writestr(it, data)

    d = docx.Document(tmp)
    # Casilla 1: AP (Additional Premium) si la prima es positiva; RP (Return Premium) si es negativa.
    transaccion = "RP" if _d(gross) < 0 else "AP"
    # Casilla 19: el % de brokerage sobre la prima (no el importe).
    brk_pct = (_d(brokerage) / _d(gross) * 100) if _d(gross) else 0
    todos = {"Premium": transaccion, "Yes/No": "Yes / No"}   # casilla 6: se deja "Yes / No"
    una_vez = {
        "Bureau": signing or "", "BrokerRef1": broker_ref1, "BrokerRef2": nombre,
        "Line": "100%", "Taxes": _num_lpan(tax), "GrossPremium": _num_lpan(gross),
        "Brokerage": _pct_lpan(brk_pct), "OCurrency": moneda, "SCurrency": moneda,
        "BureauPremium": _num_lpan(net), "UMR": umr or "",
    }
    vistos_tok: set[str] = set()
    vistos_tc: set = set()
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell._tc in vistos_tc:
                    continue
                vistos_tc.add(cell._tc)
                k = cell.text.strip()
                if k in todos:
                    _set_celda(cell, todos[k])
                elif k in una_vez:
                    _set_celda(cell, una_vez[k] if k not in vistos_tok else "")
                    vistos_tok.add(k)

    destino = os.path.join(carpeta, f"{nombre}.docx")
    d.save(destino)
    return destino


def _periodo_label(per: str) -> str:
    """'2025-01' -> 'Enero-2025'."""
    try:
        a, m = per.split("-")
        return f"{MESES_ES[int(m)]}-{a}"
    except (ValueError, IndexError):
        return per


# ──────────────────────────── Schemas ────────────────────────────
class FdoRead(BaseModel):
    id: int
    section: int
    risk_code: str
    signing_number: str | None = None
    work_package: str | None = None
    fecha_proceso: dt.date | None = None
    work_package_status: str | None = None
    fecha_generado: dt.date | None = None
    fecha_signing: dt.date | None = None
    notas: str | None = None


class LpanRead(BaseModel):
    id: int
    tipo: str
    periodo: str
    num_lineas: int
    gross_premium: Decimal | None = None
    brokerage: Decimal | None = None
    tax: Decimal | None = None
    net_premium: Decimal | None = None
    broker_ref2: str | None = None          # nombre del LPAN
    work_package: str | None = None
    signing_number: str | None = None
    fecha: dt.date | None = None            # Procesado
    sdd: dt.date | None = None
    liberado: dt.date | None = None
    pagado: dt.date | None = None
    estado: str                              # WP Status


class RiskCodeFdo(BaseModel):
    section: int
    ramo: str | None = None
    risk_code: str
    broker_reference: str       # nombre del FDO: '<parte del UMR> FDO-S<sección>-<risk code>'
    fdo: FdoRead | None = None


class RcEnSeccion(BaseModel):
    risk_code: str
    signing_number: str | None = None     # del FDO del risk code (si lo tiene)
    num_lineas: int
    gross_premium: Decimal
    brokerage: Decimal
    tax: Decimal
    net_premium: Decimal
    cobrado: bool
    lpan: LpanRead | None = None


class SeccionLpan(BaseModel):
    section: int
    risk_codes: list[RcEnSeccion]


class PeriodoLpan(BaseModel):
    periodo: str
    periodo_label: str
    secciones: list[SeccionLpan]


class VistaLpan(BaseModel):
    fdos: list[RiskCodeFdo]           # FDO/signing por risk code (transversal a periodos/secciones)
    periodos: list[PeriodoLpan]


class LpanGlobal(BaseModel):
    """Una fila del listado general de LPANs (todos los binders/pólizas)."""
    id: int
    tipo: str
    periodo: str
    binder_umr: str | None = None
    poliza_numero: str | None = None
    programa: str | None = None
    section: int
    risk_code: str
    broker_ref1: str | None = None
    broker_ref2: str | None = None
    signing_number: str | None = None
    work_package: str | None = None
    gross_premium: Decimal | None = None
    brokerage: Decimal | None = None
    tax: Decimal | None = None
    net_premium: Decimal | None = None
    fecha: dt.date | None = None       # Procesado
    sdd: dt.date | None = None
    liberado: dt.date | None = None
    pagado: dt.date | None = None
    estado: str


class FdoCreate(BaseModel):
    section: int = 0
    risk_code: str
    carpeta: str | None = None   # si se indica, genera el documento físico del FDO en esa carpeta


class FdoUpdate(BaseModel):
    signing_number: str | None = None
    work_package: str | None = None
    fecha_proceso: dt.date | None = None
    work_package_status: str | None = None
    fecha_signing: dt.date | None = None
    notas: str | None = None


class LpanCreate(BaseModel):
    risk_code: str
    section: int = 0
    periodo: str
    tipo: str = "PM"
    carpeta: str | None = None   # si se indica, genera el documento Word del LPAN en esa carpeta


class LpanUpdate(BaseModel):
    work_package: str | None = None
    fecha: dt.date | None = None       # Procesado
    sdd: dt.date | None = None
    estado: str | None = None
    liberado: dt.date | None = None
    pagado: dt.date | None = None


# ──────────────────────────── Helpers ────────────────────────────
def _binder_o_404(binder_id: int, db: Session) -> Binder:
    b = db.get(Binder, binder_id)
    if b is None:
        raise HTTPException(status_code=404, detail=f"Binder {binder_id} no encontrado")
    return b


def _secciones_declaradas(db: Session, binder_id: int) -> list[tuple[int, str | None, str]]:
    """(nº de sección, ramo, risk_code) declarados en el binder. El nº de sección es el orden
    (1-based) de las secciones, que casa con el `section_no` de los bordereaux."""
    secs = db.scalars(
        select(BinderSeccion).where(BinderSeccion.binder_id == binder_id).order_by(BinderSeccion.id)
    ).all()
    out: list[tuple[int, str | None, str]] = []
    for i, s in enumerate(secs, start=1):
        rcs = db.scalars(
            select(SeccionRiskCode).where(SeccionRiskCode.seccion_id == s.id).order_by(SeccionRiskCode.id)
        ).all()
        vistos: set[str] = set()
        for rc in rcs:
            cod = (rc.codigo or "").strip()
            if cod and cod not in vistos:
                vistos.add(cod)
                out.append((i, s.ramo, cod))
    return out


def _grupos_premium(db: Session, binder_id: int) -> dict[tuple[str, int, str], dict]:
    """Líneas del Premium (incluidas en premium) agrupadas por (periodo 'YYYY-MM', sección, risk_code)."""
    lineas = db.scalars(
        select(BdxLinea)
        .join(Bdx, BdxLinea.bdx_id == Bdx.id)
        .where(Bdx.binder_id == binder_id, BdxLinea.incluido_en_premium.is_(True), BdxLinea.premium_bdx.is_not(None))
        .options(load_only(
            BdxLinea.risk_code, BdxLinea.section_no, BdxLinea.premium_bdx,
            BdxLinea.total_gwp_our_line, BdxLinea.total_taxes_levies,
            BdxLinea.commission_coverholder_amount, BdxLinea.brokerage_amount,
            BdxLinea.final_net_premium_uw, BdxLinea.prima_cobrada,
        ))
    ).all()
    grupos: dict[tuple[str, int, str], dict] = {}
    for l in lineas:
        rc = (l.risk_code or "—").strip() or "—"
        sec = int(l.section_no) if l.section_no is not None else 0
        per = l.premium_bdx.strftime("%Y-%m")
        g = grupos.setdefault((per, sec, rc), {"num": 0, "gross": D0, "brk": D0, "tax": D0, "net": D0, "cobr": 0})
        g["num"] += 1
        g["gross"] += _d(l.total_gwp_our_line)
        g["brk"] += _d(l.commission_coverholder_amount) + _d(l.brokerage_amount)
        g["tax"] += _d(l.total_taxes_levies)
        g["net"] += _d(l.final_net_premium_uw)
        if l.prima_cobrada:
            g["cobr"] += 1
    return grupos


# ──────────────────────────── Endpoints ────────────────────────────
@router.get("/lpans", response_model=list[LpanGlobal])
def listar_lpans(db: Session = Depends(get_db)):
    """Listado GENERAL de LPANs (todos los binders y pólizas), con su contexto."""
    binders = {
        b.id: (b.umr, b.programa.nombre if b.programa else None)
        for b in db.scalars(select(Binder).options(selectinload(Binder.programa))).all()
    }
    polizas = {p.id: p.numero_poliza for p in db.scalars(select(Poliza)).all()}
    out: list[LpanGlobal] = []
    for lp in db.scalars(select(Lpan).order_by(Lpan.periodo.desc(), Lpan.id)).all():
        umr, prog = binders.get(lp.binder_id, (None, None))
        out.append(LpanGlobal(
            id=lp.id, tipo=lp.tipo, periodo=lp.periodo,
            binder_umr=umr, poliza_numero=polizas.get(lp.poliza_id), programa=prog,
            section=lp.section, risk_code=lp.risk_code,
            broker_ref1=lp.broker_ref1, broker_ref2=lp.broker_ref2,
            signing_number=lp.signing_number, work_package=lp.work_package,
            gross_premium=lp.gross_premium, brokerage=lp.brokerage, tax=lp.tax, net_premium=lp.net_premium,
            fecha=lp.fecha, sdd=lp.sdd, liberado=lp.liberado, pagado=lp.pagado, estado=lp.estado,
        ))
    return out


@router.get("/binders/{binder_id}/lpan", response_model=VistaLpan)
def vista(binder_id: int, db: Session = Depends(get_db)):
    """Vista LPAN del binder, agrupada Periodo → Sección → Risk Code, con los importes agregados,
    si está cobrado y el LPAN ya generado. Aparte, el FDO/signing por risk code (transversal)."""
    b = _binder_o_404(binder_id, db)
    grupos = _grupos_premium(db, binder_id)
    fdos = {(f.section, f.risk_code): f for f in db.scalars(select(Fdo).where(Fdo.binder_id == binder_id)).all()}
    lpans = db.scalars(select(Lpan).where(Lpan.binder_id == binder_id)).all()
    lpan_por = {(lp.periodo, lp.section, lp.risk_code): lp for lp in lpans}

    # FDO/signing por (sección, risk code) DECLARADOS en el binder (no derivados del Premium).
    # Se añaden, por si acaso, combinaciones con FDO ya creado que no estén declaradas.
    declaradas = _secciones_declaradas(db, binder_id)
    ramo_de = {(sec, rc): ramo for (sec, ramo, rc) in declaradas}
    claves = list(dict.fromkeys([(sec, rc) for (sec, _, rc) in declaradas] + sorted(fdos.keys())))
    fdos_out = [RiskCodeFdo(
        section=sec, ramo=ramo_de.get((sec, rc)), risk_code=rc,
        broker_reference=_broker_ref(b.agreement_number, sec, rc),
        fdo=FdoRead.model_validate(fdos[(sec, rc)], from_attributes=True) if (sec, rc) in fdos else None,
    ) for (sec, rc) in claves]

    periodos: list[PeriodoLpan] = []
    for per in sorted({p for (p, _, _) in grupos}, reverse=True):  # más reciente arriba
        secciones: list[SeccionLpan] = []
        for sec in sorted({s for (p, s, _) in grupos if p == per}):
            rcs: list[RcEnSeccion] = []
            for rc in sorted({r for (p, s, r) in grupos if p == per and s == sec}):
                g = grupos[(per, sec, rc)]
                lp = lpan_por.get((per, sec, rc))
                f = fdos.get((sec, rc))
                rcs.append(RcEnSeccion(
                    risk_code=rc,
                    signing_number=f.signing_number if f else None,
                    num_lineas=g["num"], gross_premium=g["gross"], brokerage=g["brk"],
                    tax=g["tax"], net_premium=g["net"],
                    cobrado=(g["num"] > 0 and g["cobr"] == g["num"]),
                    lpan=LpanRead.model_validate(lp, from_attributes=True) if lp else None,
                ))
            secciones.append(SeccionLpan(section=sec, risk_codes=rcs))
        periodos.append(PeriodoLpan(periodo=per, periodo_label=_periodo_label(per), secciones=secciones))

    return VistaLpan(fdos=fdos_out, periodos=periodos)


@router.get("/elegir-carpeta")
def elegir_carpeta(inicial: str | None = None):
    """Abre el explorador de Windows para elegir una carpeta y devuelve su ruta (solo en ejecución
    LOCAL; en servidor no hay escritorio). El diálogo se lanza en un hilo propio con su Tk."""
    import threading

    resultado: dict[str, str] = {}
    error: dict[str, str] = {}

    def _dialogo():
        try:
            import tkinter as tk
            from tkinter import filedialog
            root = tk.Tk()
            root.withdraw()
            root.attributes("-topmost", True)
            ruta = filedialog.askdirectory(
                initialdir=inicial or os.path.expanduser("~"),
                title="Elige la carpeta donde guardar el FDO",
            )
            root.destroy()
            resultado["carpeta"] = ruta or ""
        except Exception as e:  # noqa: BLE001
            error["msg"] = str(e)

    th = threading.Thread(target=_dialogo)
    th.start()
    th.join()
    if error:
        raise HTTPException(status_code=501, detail=f"No se pudo abrir el selector de carpeta: {error['msg']}")
    return {"carpeta": resultado.get("carpeta") or None}


@router.post("/binders/{binder_id}/fdo", response_model=FdoRead)
def crear_fdo(binder_id: int, payload: FdoCreate, db: Session = Depends(get_db)):
    """Genera el FDO de un risk code (a la espera del signing number de Xchanging)."""
    b = _binder_o_404(binder_id, db)
    rc = (payload.risk_code or "").strip()
    sec = int(payload.section or 0)
    if not rc:
        raise HTTPException(status_code=400, detail="Falta el risk code.")
    if db.scalar(select(Fdo).where(Fdo.binder_id == binder_id, Fdo.section == sec, Fdo.risk_code == rc)):
        raise HTTPException(status_code=409, detail=f"Ya existe un FDO para el risk code {rc} (sección {sec}).")
    # Genera el documento físico ANTES de crear el registro (si la carpeta falla, no deja huérfano).
    if payload.carpeta:
        _generar_fdo_docx(payload.carpeta, _broker_ref(b.agreement_number, sec, rc),
                          _umr_part(b.agreement_number), b.umr, None)
    f = Fdo(binder_id=binder_id, section=sec, risk_code=rc, fecha_generado=dt.date.today())
    db.add(f)
    db.commit()
    db.refresh(f)
    return FdoRead.model_validate(f, from_attributes=True)


@router.put("/fdo/{fdo_id}", response_model=FdoRead)
def actualizar_fdo(fdo_id: int, payload: FdoUpdate, db: Session = Depends(get_db)):
    """Asigna/edita el signing number (y fecha/notas) que devuelve Xchanging."""
    f = db.get(Fdo, fdo_id)
    if f is None:
        raise HTTPException(status_code=404, detail=f"FDO {fdo_id} no encontrado")
    if (f.work_package_status or "") == "Completed":
        raise HTTPException(status_code=409, detail="El FDO está en estado «Completed»: no se puede modificar.")
    datos = payload.model_dump(exclude_unset=True)
    for k, v in datos.items():
        setattr(f, k, v)
    if "signing_number" in datos and datos["signing_number"] and not f.fecha_signing:
        f.fecha_signing = dt.date.today()
    db.commit()
    db.refresh(f)
    return FdoRead.model_validate(f, from_attributes=True)


@router.delete("/fdo/{fdo_id}")
def borrar_fdo(fdo_id: int, db: Session = Depends(get_db)):
    """Borra un FDO (y, en cascada, sus LPAN)."""
    f = db.get(Fdo, fdo_id)
    if f is None:
        raise HTTPException(status_code=404, detail=f"FDO {fdo_id} no encontrado")
    db.delete(f)
    db.commit()
    return {"ok": True}


@router.post("/binders/{binder_id}/lpan", response_model=LpanRead)
def generar_lpan(binder_id: int, payload: LpanCreate, db: Session = Depends(get_db)):
    """Genera el LPAN de (risk code, periodo): exige FDO con signing y todas las líneas cobradas.
    Nombra el LPAN (Broker Ref 2), genera su documento Word (si se da carpeta) y lo deja en estado
    «Work in Progress» con WP/Procesado/SDD por rellenar."""
    b = _binder_o_404(binder_id, db)
    rc, per = (payload.risk_code or "").strip(), (payload.periodo or "").strip()
    sec = int(payload.section or 0)
    tipo = payload.tipo or "PM"
    f = db.scalar(select(Fdo).where(Fdo.binder_id == binder_id, Fdo.section == sec, Fdo.risk_code == rc))
    if f is None:
        raise HTTPException(status_code=409, detail=f"Genera antes el FDO del risk code {rc} (sección {sec}).")
    if not f.signing_number:
        raise HTTPException(status_code=409, detail=f"El FDO del risk code {rc} (sección {sec}) aún no tiene signing number.")
    grupos = _grupos_premium(db, binder_id)
    g = grupos.get((per, sec, rc))
    if not g or g["num"] == 0:
        raise HTTPException(status_code=404, detail=f"No hay líneas de Premium para {rc} (sección {sec}) en {per}.")
    if g["cobr"] != g["num"]:
        raise HTTPException(status_code=409, detail=f"Hay líneas sin cobrar en {rc} {per}: no se puede generar el LPAN.")
    if db.scalar(select(Lpan).where(Lpan.fdo_id == f.id, Lpan.periodo == per, Lpan.section == sec, Lpan.tipo == tipo)):
        raise HTTPException(status_code=409, detail=f"Ya existe un LPAN {tipo} para {rc} sección {sec} {per}.")

    nombre = _nombre_lpan(b.agreement_number, per, sec, rc)
    moneda = b.moneda or "EUR"
    # Documento Word ANTES de crear el registro (si la carpeta falla, no deja huérfano).
    if payload.carpeta:
        _generar_lpan_docx(payload.carpeta, nombre, f.signing_number, (b.agreement_number or ""),
                           b.umr, g["gross"], g["brk"], g["tax"], g["net"], moneda)
    lp = Lpan(
        fdo_id=f.id, binder_id=binder_id, risk_code=rc, section=sec, periodo=per, tipo=tipo,
        num_lineas=g["num"], gross_premium=g["gross"], brokerage=g["brk"], tax=g["tax"], net_premium=g["net"],
        broker_ref1=b.agreement_number, broker_ref2=nombre, moneda=moneda,
        work_package=None, fecha=None, sdd=None, estado="Work in Progress",
    )
    db.add(lp)
    db.commit()
    db.refresh(lp)
    return LpanRead.model_validate(lp, from_attributes=True)


@router.put("/lpan/{lpan_id}", response_model=LpanRead)
def actualizar_lpan(lpan_id: int, payload: LpanUpdate, db: Session = Depends(get_db)):
    """Edita los campos de seguimiento del LPAN: Work Package, Procesado, SDD, estado (WP Status),
    Liberado y Pagado."""
    lp = db.get(Lpan, lpan_id)
    if lp is None:
        raise HTTPException(status_code=404, detail=f"LPAN {lpan_id} no encontrado")
    for k, v in payload.model_dump(exclude_unset=True).items():
        setattr(lp, k, v)
    db.commit()
    db.refresh(lp)
    return LpanRead.model_validate(lp, from_attributes=True)


# Cabecera del Premium Bordereau (Lloyd's Coverholder Reporting Standard), tal cual el modelo.
_BDX_HEADERS = [
    "Coverholder Name", "YOA", "Unique Market Reference (UMR)", "Reporting Period Start Date",
    "Reporting Period (End Date)", "Section No", "Class of Business", "Risk Code (see list)",
    "Type of Insurance (Direct or Type or Reinsurance)", "Certificate Ref",
    "Insured Full Name, Last Name or Company Name", "ID Insured/Policyholder", "Insured Address",
    "Insured Country Sub-division: State, Province, Territory, Canton etc",
    "Insured Postcode, Zip Code or Similar", "Insured   Country (ISO code list)",
    "Risk Inception Date", "Risk Expiry Date",
    "Location of Risk - Country Sub-division: State, Province, Territory, Canton etc",
    "Location of risk - Country (ISO code)",
    "Risk, Transaction Type (New, Renewal, Endorsement, Cancellation.etc)",
    "Transaction Type - (Original premium, additional premium, return premium.etc)",
    "Cancellation Reason", "Turnover", "Effective Date of Transaction", "Expiry Date of Transaction",
    "Original Currency Premium", "Gross Premium paid this time", "Total Gross Written  Premium",
    "Fees", "Commission Coverholder %", "Commission coverholder Amount", "Total Taxes and Levies",
    "Total GWP including tax", "Net Premium to Lloyd´s Broker in original currency",
    "Sum Insured Currency (ISO code list)", "Sum Insured Amount", "Deductible Amount",
    "Deductible Basis (eec)", "Tax 1 - Jurisdiction: Country, State, Province, Territory",
    "Tax 1 - Tax Type", "Tax 1 - Amount of Taxable Premium", "Tax 1 - %", "Tax 1 - Amount",
    "Tax 1 - Administered By", "Tax 1 - Payable By",
    "Tax 2 - Jurisdiction: Country, State, Province, Territory", "Tax 2 - Tax Type",
    "Tax 2 - Amount of Taxable Premium", "Tax 2 - %", "Tax 2 - Amount", "Tax 2 - Administered By",
    "Tax 2 - Payable By", "Number of instalment", "Referred to London Yes/No", "% for Lloyd's",
    "Policy issuance date", "Policy Number Reinsured", "Brokerage % of gross premium",
    "Brokerage Amount (Original Currency)", "Final Net Premium to UW (Original Currency)",
]
# Columnas (0-based) que se subtotalizan por grupo: GWP this time, Total Taxes, Final Net to UW.
_BDX_SUBTOT = (27, 32, 60)


def _bdx_fila(l: "BdxLinea", b, coverholder: str) -> list:
    """Fila del Premium Bordereau (61 columnas) para una línea. Los % se dan como fracción (28% -> 0,28),
    igual que en el modelo. El 100% GWP y 'GWP incl. tax' no se importaron (solo our line) -> our line."""
    def f(x):
        return float(x) if x is not None else None

    def pc(x):
        return float(x) / 100 if x is not None else None    # % entero en BD -> fracción

    ourline = f(l.total_gwp_our_line)
    taxes = f(l.total_taxes_levies)
    return [
        coverholder, b.yoa, b.umr, l.reporting_period_start, l.reporting_period_end,
        l.section_no, l.class_of_business, l.risk_code, l.type_of_insurance, l.certificate_ref,
        l.insured_name, l.insured_id, l.insured_address, l.insured_province, l.insured_postcode,
        l.insured_country, l.risk_inception_date, l.risk_expiry_date, l.location_risk_province,
        l.location_risk_country, l.risk_transaction_type, l.transaction_type, None, None,
        l.effective_date_transaction, l.expiry_date_transaction, l.original_currency,
        f(l.gross_written_premium), ourline, f(l.fees), pc(l.commission_coverholder_pct),
        f(l.commission_coverholder_amount), taxes,
        (ourline or 0) + (taxes or 0), f(l.net_premium_to_broker), l.original_currency,
        f(l.sum_insured_total), f(l.deductible_amount), l.deductible_basis,
        l.tax1_jurisdiction, l.tax1_type, f(l.tax1_taxable_premium), pc(l.tax1_pct), f(l.tax1_amount),
        l.tax1_administered_by, l.tax1_payable_by,
        l.tax2_jurisdiction, l.tax2_type, f(l.tax2_taxable_premium), pc(l.tax2_pct), f(l.tax2_amount),
        l.tax2_administered_by, l.tax2_payable_by, l.number_of_instalments, l.referred_to_london,
        pc(l.pct_for_lloyds), l.policy_issuance_date, l.policy_number_reinsured, pc(l.brokerage_pct),
        f(l.brokerage_amount), f(l.final_net_premium_uw),
    ]


@router.get("/binders/{binder_id}/lpan/bdx-excel")
def bdx_excel(binder_id: int, periodo: str, db: Session = Depends(get_db)):
    """Premium Bordereau del periodo en formato Llo'yds (61 columnas), con las líneas AGRUPADAS por
    (Sección, Risk Code) como los bloques LPAN: cada grupo lleva sus filas + una fila de SUBTOTALES
    (GWP, impuestos, neto a UW) + una fila en blanco de separación."""
    import io

    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    b = _binder_o_404(binder_id, db)
    coverholder = b.productor.nombre if b.productor else (b.umr or "")
    lineas = db.scalars(
        select(BdxLinea).join(Bdx, BdxLinea.bdx_id == Bdx.id)
        .where(Bdx.binder_id == binder_id, BdxLinea.incluido_en_premium.is_(True),
               BdxLinea.premium_bdx.is_not(None))
    ).all()
    lineas = [l for l in lineas if l.premium_bdx and l.premium_bdx.strftime("%Y-%m") == periodo]

    # Agrupar por (sección, risk code), mismo orden que los bloques LPAN.
    grupos: dict[tuple, list] = {}
    for l in lineas:
        grupos.setdefault((l.section_no if l.section_no is not None else 0, (l.risk_code or "").strip()), []).append(l)

    # ── Estilo idéntico al modelo de muestra (Premium Bordereaux) ──
    ncol = len(_BDX_HEADERS)
    acc = '_-* #,##0.00_-;\\-* #,##0.00_-;_-* "-"??_-;_-@_-'    # formato contable
    acc_cols = {24, 28, 29, 30, 32, 33, 34, 35, 37, 38, 42, 44, 49, 51, 60, 61}
    pct_cols = {31, 43, 50, 56, 59}
    date_cols = {4, 5, 17, 18, 25, 26, 57, 58}
    ctr_cols = {2, 4, 5, 17, 18, 25, 26}
    rgt_cols = {38, 57}
    widths = {3: 18.7, 10: 20.9, 11: 20.7, 12: 12.0, 13: 19.6, 24: 12.0}   # resto: 13
    sub_1b = {i + 1 for i in _BDX_SUBTOT}     # columnas (1-based) con subtotal: 28, 33, 61

    def numfmt(col: int) -> str:
        if col in acc_cols:
            return acc
        if col in pct_cols:
            return "0.00%"
        if col in date_cols:
            return "mm-dd-yy"
        if col == 15:
            return "00000"
        return "General"

    thin = Side(style="thin")
    hdr_font = Font(name="Calibri", size=9, bold=True)
    hdr_fill = PatternFill("solid", fgColor="D9D9D9")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    hdr_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    data_font = Font(name="Aptos Narrow", size=9)
    sub_font = Font(name="Aptos Narrow", size=9, bold=True)
    sub_border = Border(top=thin, bottom=thin)
    ctr, rgt = Alignment(horizontal="center"), Alignment(horizontal="right")

    def estilo_fila(row: int, subtotal: bool = False) -> None:
        for col in range(1, ncol + 1):
            c = ws.cell(row=row, column=col)
            c.font = sub_font if subtotal else data_font
            c.number_format = numfmt(col)
            if subtotal:
                if col in sub_1b:
                    c.border = sub_border
            elif col in ctr_cols:
                c.alignment = ctr
            elif col in rgt_cols:
                c.alignment = rgt

    meses_en = ["", "January", "February", "March", "April", "May", "June",
                "July", "August", "September", "October", "November", "December"]
    try:
        titulo_hoja = meses_en[int(periodo.split("-")[1])]
    except (ValueError, IndexError):
        titulo_hoja = periodo

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = titulo_hoja
    ws.append(_BDX_HEADERS)
    for col in range(1, ncol + 1):
        c = ws.cell(row=1, column=col)
        c.font, c.fill, c.alignment, c.border = hdr_font, hdr_fill, hdr_align, hdr_border
    ws.row_dimensions[1].height = 96

    for clave in sorted(grupos.keys()):
        filas = grupos[clave]
        for l in filas:
            ws.append(_bdx_fila(l, b, coverholder))
            estilo_fila(ws.max_row)
        # Subtotales del grupo (solo en las 3 columnas de importe).
        sub = [None] * ncol
        for idx, campo in zip(_BDX_SUBTOT, ("gross_written_premium", "total_taxes_levies", "final_net_premium_uw")):
            sub[idx] = float(sum((getattr(l, campo) or 0) for l in filas))
        ws.append(sub)
        estilo_fila(ws.max_row, subtotal=True)
        ws.append([])   # separación entre grupos

    for col in range(1, ncol + 1):
        ws.column_dimensions[get_column_letter(col)].width = widths.get(col, 13.0)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"Premium Bordereaux {b.umr or b.agreement_number or binder_id} {periodo}.xlsx"
    return StreamingResponse(
        buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.delete("/lpan/{lpan_id}")
def borrar_lpan(lpan_id: int, db: Session = Depends(get_db)):
    lp = db.get(Lpan, lpan_id)
    if lp is None:
        raise HTTPException(status_code=404, detail=f"LPAN {lpan_id} no encontrado")
    db.delete(lp)
    db.commit()
    return {"ok": True}
